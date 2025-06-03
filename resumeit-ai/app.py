from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import tempfile
import fitz  
import docx
import magic
from sentence_transformers import SentenceTransformer, util
from dotenv import load_dotenv
load_dotenv()


# Initialize Flask app
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# Load embedding model for resume scoring
model = SentenceTransformer('all-MiniLM-L6-v2')

# === Text Extraction ===
def extract_text(file_path):
    try:
        mime = magic.Magic(mime=True)
        mime_type = mime.from_file(file_path)

        if mime_type == "application/pdf":
            with fitz.open(file_path) as doc:
                return "".join([page.get_text() for page in doc]).strip() or "[NO_TEXT_FOUND_PDF]"

        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs]).strip() or "[NO_TEXT_FOUND_DOCX]"

        elif mime_type == "text/plain":
            with open(file_path, 'r', encoding="utf-8") as f:
                return f.read().strip() or "[NO_TEXT_FOUND_TXT]"

        return f"[UNSUPPORTED_MIME: {mime_type}]"

    except Exception as e:
        return f"[EXTRACTION_FAILED: {str(e)}]"


# === Text Cleaning ===
def clean_text(text):
    return text.replace('\n', ' ').replace('\r', ' ').strip()


def truncate_text(text, max_chars=800):
    return text[:max_chars] + "..." if len(text) > max_chars else text

from google.cloud import language_v1

def extract_entities(text):
    client = language_v1.LanguageServiceClient()
    document = language_v1.Document(content=text, type_=language_v1.Document.Type.PLAIN_TEXT)
    response = client.analyze_entities(document=document)

    entities = [entity.name for entity in response.entities if entity.salience > 0.01]
    return list(set(entities))  # remove duplicates
import random

def generate_interview_questions(keywords, count=5):
    templates = [
        "Can you explain your experience with {}?",
        "How have you applied {} in your previous role?",
        "What challenges have you faced while working with {}?",
        "Describe a project where you used {}.",
        "How would you improve efficiency using {}?"
    ]
    questions = []

    for keyword in random.sample(keywords, min(count, len(keywords))):
        template = random.choice(templates)
        questions.append(template.format(keyword))

    return questions

@app.route('/generate-questions', methods=['POST'])
def generate_questions():
    job_file = request.files.get("job_description")
    resume_file = request.files.get("resume")

    if not job_file or not resume_file:
        return jsonify({"error": "Both job_description and resume files are required"}), 400

    with tempfile.NamedTemporaryFile(delete=False) as jd_temp:
        job_file.save(jd_temp.name)
        jd_text = clean_text(extract_text(jd_temp.name))
        os.unlink(jd_temp.name)

    with tempfile.NamedTemporaryFile(delete=False) as res_temp:
        resume_file.save(res_temp.name)
        resume_text = clean_text(extract_text(res_temp.name))
        os.unlink(res_temp.name)

    combined_text = f"{jd_text} {resume_text}"
    keywords = extract_entities(combined_text)

    if not keywords:
        return jsonify({"error": "Could not extract relevant entities from text"}), 400

    questions = generate_interview_questions(keywords)
    return jsonify({"questions": questions})

# === Endpoint: Score Resumes ===
@app.route('/score-resumes', methods=['POST'])
def score_resumes():
    job_file = request.files.get("job_description")
    resumes = request.files.getlist("resumes")

    if not job_file or not resumes:
        return jsonify({"error": "Both job_description and resume files are required"}), 400

    with tempfile.NamedTemporaryFile(delete=False) as temp_jd:
        job_file.save(temp_jd.name)
        job_desc_text = clean_text(extract_text(temp_jd.name))
        os.unlink(temp_jd.name)

    jd_embedding = model.encode(job_desc_text, convert_to_tensor=True)
    results = []

    for resume in resumes:
        with tempfile.NamedTemporaryFile(delete=False) as temp_resume:
            resume.save(temp_resume.name)
            resume_text = clean_text(extract_text(temp_resume.name))
            os.unlink(temp_resume.name)

        resume_embedding = model.encode(resume_text, convert_to_tensor=True)
        similarity_score = util.cos_sim(jd_embedding, resume_embedding).item()

        results.append({
            "filename": resume.filename,
            "score": round(similarity_score * 100, 2)
        })

    return jsonify({"results": results})


# === App Entry Point ===
if __name__ == '__main__':
    app.run(port=8000, debug=True)